"""
resume_parser/extractor.py — Raw Text Extractor
===============================================
Purpose
-------
Extract raw text from PDF, DOCX, or TXT resume files.

Design Decisions
----------------
Why pdfplumber over other PDF libraries (e.g. PyPDF2, pypdf)?
    - pdfplumber preserves document layout and column ordering. Resumes are
      frequently designed with multi-column formats. Standard parsers often
      mix text across columns, rendering section detection impossible.
    - pdfplumber provides detailed control over character coordinates and spacing,
      resulting in much higher parsing fidelity.

Why python-docx?
    - Simple, robust library for reading Microsoft Word files (.docx).
    - Extracts paragraphs, tables, and headers natively without requiring
      external office suites or platforms.

Fallback to UTF-8 / CP1252:
    - Plain text files (.txt) can be encoded in various formats.
    - We try UTF-8 first (the industry standard) and fall back to CP1252 (the
      default Windows encoding) to prevent UnicodeDecodeError.

Usage
-----
    from resume_parser.extractor import ResumeExtractor
    from pathlib import Path

    extractor = ResumeExtractor()
    raw_text = extractor.extract(Path("resume/my_resume.pdf"))
    print(raw_text[:200])
"""

from __future__ import annotations

from pathlib import Path

import pdfplumber
import docx

from utils.exceptions import ResumeParserError
from utils.logger import get_logger

logger = get_logger(__name__)


class ResumeExtractor:
    """
    Extracts plain text from resume files depending on their file extension.
    """

    def extract(self, file_path: Path) -> str:
        """
        Extract text from the specified file path.

        Parameters
        ----------
        file_path : Path
            Path to the PDF, DOCX, or TXT file.

        Returns
        -------
        str
            Extracted raw text.

        Raises
        ------
        ResumeParserError
            If file reading or parsing fails, or format is unsupported.
        """
        path = Path(file_path)
        if not path.exists():
            raise ResumeParserError(f"Resume file not found: {path}", path=str(path))

        ext = path.suffix.lower()

        logger.info("Extracting text from resume", extra={"path": str(path), "ext": ext})

        try:
            if ext == ".pdf":
                return self._extract_pdf(path)
            elif ext == ".docx":
                return self._extract_docx(path)
            elif ext == ".txt":
                return self._extract_txt(path)
            else:
                raise ResumeParserError(
                    f"Unsupported format: {ext}. Only PDF, DOCX, and TXT are supported.",
                    path=str(path)
                )
        except ResumeParserError:
            raise
        except Exception as exc:
            logger.error(
                "Failed to extract text from resume file",
                extra={"path": str(path), "error": str(exc)}
            )
            raise ResumeParserError(
                f"Failed to extract text from '{path.name}': {exc}",
                path=str(path)
            ) from exc

    # -------------------------------------------------------------------------
    # Format-specific extractors
    # -------------------------------------------------------------------------

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from a PDF file using pdfplumber."""
        pages_text = []
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    # extract_text returns None if the page has no text (e.g. image-only)
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                    else:
                        logger.warning(
                            "PDF page contains no extractable text (might be scanned/image)",
                            extra={"path": str(path), "page": i}
                        )
        except Exception as exc:
            raise ResumeParserError(
                f"Error parsing PDF file with pdfplumber: {exc}",
                path=str(path)
            ) from exc

        full_text = "\n\n".join(pages_text)
        if not full_text.strip():
            raise ResumeParserError(
                "Extracted text is empty. The PDF may be scanned or image-only.",
                path=str(path)
            )

        return full_text

    def _extract_docx(self, path: Path) -> str:
        """Extract text from a Word DOCX file using python-docx."""
        try:
            doc = docx.Document(path)
            # 1. Extract normal paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # 2. Extract table cell contents (often used for Skills/Experience layout)
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        tables_text.append(" | ".join(row_cells))

            # Combine them all
            full_text = "\n\n".join(paragraphs + tables_text)
            if not full_text.strip():
                raise ResumeParserError(
                    "Extracted text from DOCX is empty.",
                    path=str(path)
                )
            return full_text
        except Exception as exc:
            raise ResumeParserError(
                f"Error parsing DOCX file: {exc}",
                path=str(path)
            ) from exc

    def _extract_txt(self, path: Path) -> str:
        """Extract text from a plain TXT file trying UTF-8 first, then CP1252."""
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            logger.warning(
                "Failed to decode TXT file as UTF-8. Retrying with CP1252.",
                extra={"path": str(path)}
            )
            try:
                return path.read_text(encoding="cp1252")
            except Exception as exc:
                raise ResumeParserError(
                    f"Error reading plain text file: {exc}",
                    path=str(path)
                ) from exc
        except Exception as exc:
            raise ResumeParserError(
                f"Error reading plain text file: {exc}",
                path=str(path)
            ) from exc
