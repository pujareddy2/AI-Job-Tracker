"""
communication_engine/exporter.py — Document Export System
==========================================================
Purpose
-------
Exports generated documents to TXT, Markdown, HTML, DOCX, and PDF.

Design Decisions
----------------
- Standard text file writing for TXT and Markdown.
- HTML styled responsive container.
- python-docx Paragraph creation for DOCX.
- reportlab Platypus Document Template flowable rendering for PDF.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT

from utils.logger import get_logger

logger = get_logger(__name__)


class DocumentExporter:
    """Handles multi-format exports of outreach documents."""

    @staticmethod
    def export(
        doc_type: str,
        tone: str,
        subject: str | None,
        body: str,
        export_dir: Path,
        formats: list[str],
    ) -> dict[str, str]:
        """
        Export a single document to the specified formats.

        Returns a dictionary mapping format to the absolute file path.
        """
        export_dir.mkdir(parents=True, exist_ok=True)
        safe_name = doc_type.lower().replace(" ", "_").replace("/", "_")
        safe_tone = tone.lower()
        base_filename = f"{safe_name}_{safe_tone}"
        
        exported_paths = {}

        for fmt in formats:
            fmt_lower = fmt.lower().strip(".")
            if fmt_lower == "txt":
                path = DocumentExporter._export_txt(body, subject, export_dir / f"{base_filename}.txt")
                exported_paths["txt"] = str(path)
            elif fmt_lower == "md":
                path = DocumentExporter._export_md(doc_type, tone, body, subject, export_dir / f"{base_filename}.md")
                exported_paths["md"] = str(path)
            elif fmt_lower == "html":
                path = DocumentExporter._export_html(doc_type, tone, body, subject, export_dir / f"{base_filename}.html")
                exported_paths["html"] = str(path)
            elif fmt_lower == "docx":
                path = DocumentExporter._export_docx(doc_type, tone, body, subject, export_dir / f"{base_filename}.docx")
                exported_paths["docx"] = str(path)
            elif fmt_lower == "pdf":
                path = DocumentExporter._export_pdf(doc_type, tone, body, subject, export_dir / f"{base_filename}.pdf")
                exported_paths["pdf"] = str(path)

        return exported_paths

    # ── Exporters for each format ────────────────────────────────────────────

    @staticmethod
    def _export_txt(body: str, subject: str | None, dest: Path) -> Path:
        content = ""
        if subject:
            content += f"Subject: {subject}\n\n"
        content += body
        dest.write_text(content, encoding="utf-8")
        return dest

    @staticmethod
    def _export_md(doc_type: str, tone: str, body: str, subject: str | None, dest: Path) -> Path:
        content = f"# {doc_type} ({tone} Version)\n\n"
        if subject:
            content += f"**Subject:** {subject}\n\n"
        content += body
        dest.write_text(content, encoding="utf-8")
        return dest

    @staticmethod
    def _export_html(doc_type: str, tone: str, body: str, subject: str | None, dest: Path) -> Path:
        body_html = body.replace("\n", "<br>")
        subject_html = f"<div class='subject'><strong>Subject:</strong> {subject}</div>" if subject else ""
        
        html_template = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<title>{doc_type}</title>
<style>
    body {{
        font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        background-color: #f4f7f6;
        color: #333;
        margin: 0;
        padding: 40px 20px;
    }}
    .container {{
        max-width: 600px;
        background: #ffffff;
        margin: 0 auto;
        padding: 40px;
        border-radius: 12px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.05);
        border: 1px solid #e1e8ed;
    }}
    h1 {{
        font-size: 24px;
        color: #0f172a;
        margin-top: 0;
        border-bottom: 2px solid #3b82f6;
        padding-bottom: 12px;
    }}
    .meta {{
        font-size: 14px;
        color: #64748b;
        margin-bottom: 20px;
    }}
    .subject {{
        background: #f1f5f9;
        padding: 12px;
        border-radius: 6px;
        margin-bottom: 20px;
        font-size: 15px;
    }}
    .body {{
        font-size: 16px;
        line-height: 1.6;
        color: #334155;
    }}
</style>
</head>
<body>
    <div class="container">
        <h1>{doc_type}</h1>
        <div class="meta">Tone: {tone}</div>
        {subject_html}
        <div class="body">
            {body_html}
        </div>
    </div>
</body>
</html>
"""
        dest.write_text(html_template, encoding="utf-8")
        return dest

    @staticmethod
    def _export_docx(doc_type: str, tone: str, body: str, subject: str | None, dest: Path) -> Path:
        docx = DocxDocument()
        docx.add_heading(f"{doc_type} ({tone} Version)", level=1)
        if subject:
            docx.add_paragraph().add_run(f"Subject: {subject}").bold = True
        
        # Add paragraphs
        for p in body.split("\n\n"):
            if p.strip():
                docx.add_paragraph(p.strip())
                
        docx.save(dest)
        return dest

    @staticmethod
    def _export_pdf(doc_type: str, tone: str, body: str, subject: str | None, dest: Path) -> Path:
        # Create a reportlab SimpleDocTemplate
        doc = SimpleDocTemplate(
            str(dest),
            pagesize=letter,
            rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        
        # Define clean styles
        title_style = ParagraphStyle(
            name="TitleStyle",
            parent=styles["Heading1"],
            fontSize=22,
            leading=26,
            textColor="#0f172a",
            spaceAfter=15,
            alignment=TA_LEFT
        )
        
        meta_style = ParagraphStyle(
            name="MetaStyle",
            fontSize=11,
            leading=14,
            textColor="#64748b",
            spaceAfter=20
        )
        
        subject_style = ParagraphStyle(
            name="SubjectStyle",
            fontSize=12,
            leading=16,
            textColor="#1e293b",
            backColor="#f1f5f9",
            borderColor="#cbd5e1",
            borderWidth=1,
            borderPadding=10,
            spaceAfter=20
        )
        
        body_style = ParagraphStyle(
            name="BodyStyle",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            textColor="#334155",
            spaceAfter=12,
            alignment=TA_JUSTIFY
        )

        flowables = []
        flowables.append(Paragraph(f"{doc_type}", title_style))
        flowables.append(Paragraph(f"Tone: {tone}", meta_style))
        
        if subject:
            flowables.append(Paragraph(f"<b>Subject:</b> {subject}", subject_style))
            flowables.append(Spacer(1, 10))
            
        # Parse paragraphs
        for p in body.split("\n\n"):
            if p.strip():
                # Replace newlines inside paragraph with break tag
                cleaned_p = p.strip().replace("\n", "<br/>")
                flowables.append(Paragraph(cleaned_p, body_style))
                
        doc.build(flowables)
        return dest
